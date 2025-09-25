from docx.api import Document
from docx.opc.exceptions import PackageNotFoundError
import document as doc
from pathlib import Path
from core.resolution import *
import re
import roman
from typing import Generic, TypeVar, cast, Callable
import json
from colorama import Fore, Back, init, Style
import argparse
import os
init() # colorama

verbose: bool = False

# ==== CONFIG ====

preamb_config_path = Path("./config/preambs/config.json")
operationals_config_path = Path("./config/operationals/config.json")

preamb_config: dict[str, list[str] | str | dict] = {}
with preamb_config_path.open("r", encoding="utf-8") as f:
    preamb_config = json.load(f)

operationals_config: dict[str, list[str] | str | dict] = {}
with operationals_config_path.open("r", encoding="utf-8") as f:
    operationals_config = json.load(f)

preamb_phrases = sorted(preamb_config.get('preambs_phrases', []))
operationals_phrases = sorted(cast(list[str], operationals_config.get('operationals_phrases', [])))
list_phrases = sorted(operationals_config.get("list_phrases", []))

if preamb_phrases == []:
    print(f"{Fore.RED}Warning: no preamb phrases loaded{Style.RESET_ALL}")
if operationals_phrases == []:
    print(f"{Fore.RED}Warning: no operational phrases loaded{Style.RESET_ALL}")
if list_phrases == []:
    print(f"{Fore.RED}Warning: no list phrases loaded{Style.RESET_ALL}")

if verbose: print(f"{Fore.GREEN}{len(preamb_phrases)} preamb phrases loaded.{Style.RESET_ALL}")
if verbose: print(f"{Fore.GREEN}{len(operationals_phrases)} operational phrases loaded.{Style.RESET_ALL}")
if verbose: print(f"{Fore.GREEN}{len(list_phrases)} list phrases loaded.{Style.RESET_ALL}")
# ====

# print("Loading language package. This may take a while.")
# import spacy
# nlp = spacy.load('en_core_web_sm')

class ResolutionParsingError(BaseException):
    def __init__(self, msg: str, line: int):
        super().__init__(msg)
        self.line = line
    def __str__(self):
        return f"LINE {self.line}: {super().__str__()}"
    
def strip_punctuations(text: str) -> str:
    """
    Removes extra space for punctuations when tokens of a sentence is joined by a space literal
    """
    for punc in [',', '.', ';', ':', '!', '?']:
        text = text.replace(f" {punc}", f"{punc}")
    return text

"""
def extract_first_participial_phrase(text: str) -> tuple[str | None, str]:
    doc = nlp(text)
    
    for i, token in enumerate(doc):
        # Check for participles more carefully
        is_participle = (
            token.pos_ == 'VERB' and 
            token.morph.get('VerbForm', None) == ['Part']
        )
        
        # Also check for words ending with -ing, -ed, -en that might be participles
        is_likely_participle = (
            token.pos_ == 'VERB' and 
            (token.text.endswith('ing') or 
             token.text.endswith('ed') or 
             token.text.endswith('en'))
        )
        
        if is_participle or is_likely_participle:
            # Start building the participial phrase
            phrase_tokens = [token.text]
            j = i + 1
            
            # Include modifiers, prepositions, particles, etc.
            while (j < len(doc) and 
                  (doc[j].pos_ in ['ADP', 'ADV', 'PART', 'SCONJ'] or 
                   doc[j].dep_ in ['prep', 'agent', 'pcomp', 'mark', 'advmod'])):
                phrase_tokens.append(doc[j].text)
                j += 1
            
            # Extract the first participial phrase
            participial_phrase = ' '.join(phrase_tokens)
            remaining_text = ' '.join([t.text for t in doc[j:]])
            remaining_text = strip_punctuations(remaining_text)
            print("Participial phrase: " + participial_phrase + ", remaining: " + remaining_text)
            return participial_phrase, remaining_text
    
    return None, text

# Doesn't work for 'be' verbs and 'will' aux verbs
def extract_first_verb(text: str) -> tuple[str | None, str | None, bool]:
    \"""
    Extracts first verb from a sentence. Includes a preposition after the verb if there is any (used for the clause verb extraction). Returns the verb and the preposition, the rest of the sentence, and whether it occurs at the beginning of the sentence.
    Params:
    ------------
    text: str
        The text in which to extract the first verb
    Returns:
    ------------
    tuple [str | None, str, bool]
        1. The first element (str | None): 
            the verb of the sentence. If there is no verb in the sentence or if the sentence is empty, returns None

        2. The second element (str | None): 
            the rest of the sentence. If the sentence is empty, returns None

        3. The third element (bool):
            whether the verb occurs at the beginning of the sentence. If it does not, or if the text is empty, or if the text contains no verb, returns False
    \"""
    if not text or text.isspace():
        return None, None, False

    doc = nlp(text)
    
    for i, token in enumerate(doc):
        if token.pos_ in ['VERB', 'AUX']:
            # Check if this is part of a phrasal verb
            verb_phrase = token.text
            
            # Look for particles/prepositions that might be part of phrasal verbs
            if i + 1 < len(doc) and doc[i+1].dep_ in ['prt', 'prep']:
                verb_phrase += " " + doc[i+1].text
            
            # Reconstruct the rest of the sentence
            tokens_to_remove = [token]
            if i + 1 < len(doc) and doc[i+1].dep_ in ['prt', 'prep']:
                tokens_to_remove.append(doc[i+1])
            
            rest_of_sentence = ' '.join([
                t.text for t in doc 
                if t not in tokens_to_remove and not t.is_space
            ]).strip()
            
            # Clean up punctuation
            rest_of_sentence = strip_punctuations(rest_of_sentence)
            return verb_phrase, rest_of_sentence, i == 0
    
    # No verb found
    rest_of_sentence = ' '.join([t.text for t in doc if not t.is_space]).strip()
    rest_of_sentence = strip_punctuations(rest_of_sentence)
    return None, rest_of_sentence, False
"""

# === TYPES to silence the type checker ===
type _rc_inner_t = str | preamb | clause
type _rc_t = ResolutionComponent[_rc_inner_t]
T = TypeVar('T')
type _mat_func_t[T] = Callable[[str], tuple[T | None, bool]]

class ResolutionComponent(Generic[T]):
    def __init__(self, startIdx: int = -1, endIdx: int = -1, patterns: list[str] | None = None, currentIdx: int = 0, matchFunc: _mat_func_t[T] | None = None):
        self.startIdx = startIdx
        self.endIdx = endIdx
        self.patterns = patterns if patterns is not None else [r'(.*)']
        self.currentIdx = currentIdx
        self.parsed = False
        self.values: list[T] = [] # single element list if not a list
        self.results: list[re.Match | None] | None = None
        self.found = False
        self.matchFunc: _mat_func_t[T] | None = cast(_mat_func_t[T], matchFunc)

    
    def extract(self, text: str, flag: int | None = None) -> None:
        """Extract content using the pattern and store result"""
        # if self.parsed: return
        if self.matchFunc is None:
            self.results = self.getContentFrom(text, flag)
            if self.results is not None:
                if any(self.results):
                    for result in self.results:
                        if result is not None:
                            extracted_value = result.group(1).strip() if result.groups() else result.group(0).strip()
                            self.appendValue(extracted_value)
        else:
            # matchFunc returns tuple[str, bool]
            # str: remaining, bool: format correctness
            result =self.matchFunc(text)
            if (result[1]):
                self.appendValue(cast(T, result[0]))

    def markFinished(self):
        self.parsed = True
    
    def getContentFrom(self, text: str, flag: int | None = None) -> list[re.Match | None] | None: # None here because it returns early if its parsed already
        # if self.parsed: return
        searchList: list[re.Match | None] = []
        if flag is not None:
            for pattern in self.patterns:
                searchList.append(re.search(pattern, text, flag))
            return searchList
        for pattern in self.patterns:
            searchList.append(re.search(pattern, text))
        return searchList
    
    def setValue(self, values: list[T]) -> None:
        self.values = values
    
    def appendValue(self, newValue: T) -> None:
        self.values.append(newValue)
    
    def getValues(self) -> list[T]:
        return self.values
    
    def getFirst(self) -> T | str:
        if self.values:
            return self.values[0]
        return ""
        
    
    def getListValues(self, delimiter: str = ",") -> list[T] | None:
        """Returns a list of values separated by a delimiter if the type is str"""
        if self.values and not all(isinstance(val, str) for val in self.values):
            return None
        
        if not self.getValues():
            return []
        
        if len(self.getValues()) > 1:
            return self.getValues()
        
        first = self.getFirst()
        if first is not None and isinstance(first, str) and delimiter in first:
            return [item.strip() for item in first.split(delimiter)] # type: ignore
        return self.getValues()

def parseToResolution (doc: doc.document)\
      -> tuple[Resolution, dict[str, _rc_t], list[ResolutionParsingError]]:
    # TODO: implement security council formatting
    paragraphs = doc.get_paragraphs()

    components: dict[str, ResolutionComponent[_rc_inner_t]] = {}
    errorList: list[ResolutionParsingError] = []

    # make patterns more permissive (allow colon, dash, or whitespace separators)
    components['committee'] = cast(_rc_t, ResolutionComponent[str](patterns=[
        r'committee\s*[:\-\s]\s*(.*)', r'committee[:\-\s]*(.*)',
    ]))

    components['mainSubmitter'] = cast(_rc_t, ResolutionComponent[str](patterns=[
        r'main[\s\-]*submitter[s]?\s*[:\-\s]\s*(.*)',
        r'main[\s\-]*submitter[s]?\s*[:\-\s]*(.*)',
    ]))

    components['coSubmitters'] = cast(_rc_t, ResolutionComponent[str](patterns=[
        r'co[\s\-]*submitter[s]?\s*[:\-\s]\s*(.*)', r'co[\s\-]*submitter[s]?\s*[:\-\s]*(.*)',
        r'cosubmitter[s]?\s*[:\-\s]\s*(.*)', r'cosubmitter[s]?\s*[:\-\s]*(.*)',
    ]))

    components['topic'] = cast(_rc_t, ResolutionComponent[str](patterns=[
        r'topic[s]?\s*[:\-\s]\s*(.*)', r'topic[s]?\s*[:\-\s]*(.*)',
    ]))

    listPreambs: list[preamb] = []
    listOperationals: list[clause] = []

    def sanitize_text(s: str) -> str:
        """Trim whitespace, collapse duplicate punctuation, strip trailing commas/semicolons/colons and extra spaces."""
        if s is None:
            return ""
        s = s.strip()
        # replace multiple commas or semicolons with single
        while ',,' in s:
            s = s.replace(',,', ',')
        while ';;' in s:
            s = s.replace(';;', ';')
        # remove trailing punctuation that should not remain
        s = s.rstrip(' ,;:')
        # collapse multiple spaces
        s = re.sub(r'\s+', ' ', s)
        return s
    
    def normalize_committee(name: str) -> str:
        name = re.sub(r'\s*\([^)]*\)', '', name)
        return name.strip().lower()

    def is_intro_line(text: str, committee_name: str) -> bool:
        norm_text = re.sub(r'[,\s]+$', '', text.strip()).lower()
        norm_committee = normalize_committee(committee_name)
        return norm_text == f"the {norm_committee}"


    def _preambs_match_function(text: str) -> tuple[preamb | None, bool]:
        """
        Parse one line of preambular text.

        Returns:
        (preamb_obj, True)  -> created a valid preambular clause
        (preamb_obj, False) -> could not match
        """
        raw = text.strip()
        if not raw:
            return (preamb("__EMPTY__", ""), False)
        if is_intro_line(raw, cast(str, components['committee'].getValues()[0])):
            return None, False
        # If the line appears to be a numbered / lettered / roman operational line,
        # don't try to treat it as a preamb.
        if re.match(r'^\s*(\d+[\.\)]|\(?[A-Za-z][\.\)]|\(?[ivxIVX]+\s*[\.\)])', raw):
            return (preamb("__SKIP__", raw), False)

        # Try longest phrases first so we don't prematurely match a short phrase that's
        # a substring of a longer one.
        for phrase in sorted(set(preamb_phrases), key=len, reverse=True):
            ph_esc = re.escape(phrase)
            # match phrase at start (optionally preceded by quotes/parentheses) and capture remainder
            pattern = rf'^\s*[\(\["\']*\s*{ph_esc}\b(?:\s+(.*?))?\s*[,;]?\s*$'
            m = re.match(pattern, raw, flags=re.IGNORECASE)
            if m:
                remainder = (m.group(1) or "").strip()
                remainder = sanitize_text(strip_punctuations(remainder))
                p = preamb(phrase, remainder)
                listPreambs.append(p)
                return (p, True)

        # Header-like lines that end with comma/colon and are not numbered -> treat as preamb header
        if raw.endswith(',') or raw.endswith(':'):
            head = raw.rstrip(',;:').strip()
            # ensure it's not a numbered/lettered/roman header
            if not re.match(r'^\s*(\d+[\.\)]|\(?[A-Za-z][\.\)]|\(?[ivxIVX]+\s*[\.\)])', head):
                p = preamb(head, "")
                listPreambs.append(p)
                return (p, True)

        # No match
        return (preamb("__ERROR__", raw), False)

    def dedupe_preserve_order(seq):
        seen = set()
        out = []
        for s in seq:
            if s not in seen:
                seen.add(s)
                out.append(s)
        return out

    def _find_first_operational_phrase(text: str, phrases: list[str]) -> tuple[str | None, str, bool]:
        if not phrases:
            return (None, text, False)
        sorted_phrases = sorted(set(phrases), key=lambda s: len(s), reverse=True)
        orig = text
        # first try to match phrase at start (allow leading punctuation/quotes)
        for ph in sorted_phrases:
            ph_esc = re.escape(ph)
            pattern_start = rf'^\s*[\(\["\']*\s*({ph_esc})\b'
            m = re.search(pattern_start, orig, flags=re.IGNORECASE)
            if m:
                start_idx = m.start(1)
                end_idx = m.end(1)
                rest = orig[end_idx:].lstrip(" \t\n\r:;,-—–.()[]\"'")
                rest = sanitize_text(strip_punctuations(rest))
                return (orig[start_idx:end_idx], rest, True)

        # If no start matches, search for the first occurrence but ensure it's not part of a
        # numbered/lettered prefix (so we don't pick up internal words from subclauses accidentally).
        for ph in sorted_phrases:
            ph_esc = re.escape(ph)
            pattern_any = rf'\b({ph_esc})\b'
            m = re.search(pattern_any, orig, flags=re.IGNORECASE)
            if m:
                # ensure the match isn't inside a leading clause marker like "1." or "a)"
                prefix = orig[:m.start(1)]
                if re.search(r'^\s*$', prefix) or not re.search(r'[\dA-Za-z\)\.]\s*$', prefix):
                    rest = orig[m.end(1):].lstrip(" \t\n\r:;,-—–.()[]\"'")
                    rest = sanitize_text(strip_punctuations(rest))
                    return (orig[m.start(1):m.end(1)], rest, False)
        return (None, text, False)

    def _operationals_match_function(text: str) -> tuple[clause, bool]:
        global verbose

        # persistent state stored on the function object
        if not hasattr(_operationals_match_function, "state"):
            _operationals_match_function.state = { # type: ignore
                "clause_counter"        : 0,
                "subclause_counter"     : 0,
                "subsubclause_counter"  : 0,
                "current_clause"        : None,
                "current_subclause"     : None,
                "current_subsubclause"  : None,
            }
        st = _operationals_match_function.state # type: ignore

        raw = text.strip()
        if not raw:
            return (st["current_clause"] if st["current_clause"] is not None else clause(0, "__EMPTY__", "__EMPTY__"), False)

        def roman_to_int_lower(s: str) -> int | None:
            try:
                return roman.fromRoman(s.upper())
            except Exception:
                return None

        # Top-level clause detection (Arabic numerals)
        m_clause = re.match(r'^\s*(\d+)\s*[\.\)]\s*(.+)$', raw)
        if m_clause:
            num = m_clause.group(1)
            body = m_clause.group(2).strip()
            try:
                idx = int(num)
            except Exception:
                st["clause_counter"] += 1
                idx = st["clause_counter"]

            st["clause_counter"] = max(st["clause_counter"], idx)

            # detect verb phrase using operationals_phrases (if available)
            try:
                phrases = operationals_phrases
            except NameError:
                phrases = []

            verb_phrase, rest_of_sentence, at_start = _find_first_operational_phrase(body, phrases)
            if verb_phrase:
                verb = sanitize_text(verb_phrase.strip())
                clause_text = sanitize_text(rest_of_sentence if rest_of_sentence else "")
            else:
                verb = "clause verb"
                clause_text = sanitize_text(body)

            new_clause = clause(idx, verb=verb, text=clause_text)
            st["current_clause"      ] = new_clause
            st["current_subclause"   ] = None
            st["current_subsubclause"] = None
            st["subclause_counter"   ] = 0
            st["subsubclause_counter"] = 0

            try:
                listOperationals.append(new_clause)
            except NameError:
                pass

            return (new_clause, True)

        # Sub-subclause detection: roman numerals (i, ii, iii, ...)
        m_ssc = re.match(r'^\s*\(?([ivxIVX]{1,4})\)?\s*[\.\)]\s*(.+)$', raw)
        if m_ssc:
            roman_str = m_ssc.group(1)
            body = sanitize_text(m_ssc.group(2).strip())
            roman_idx = roman_to_int_lower(roman_str)
            if st["current_subclause"] is not None:
                idx = roman_idx if roman_idx is not None else (st["subsubclause_counter"] + 1)
                st["subsubclause_counter"] += 1
                new_ssc = subsubclause(idx, text=body)
                st["current_subclause"   ].append(new_ssc)
                st["current_subsubclause"] = new_ssc
                return (st["current_clause"], False)
            # fallback: create a new subclause and attach this as its first sub-subclause
            if st["current_clause"] is not None:
                st["subclause_counter"] += 1
                new_sub = subclause(st["subclause_counter"], text=body)
                st["current_clause"      ].listsubclauses.append(new_sub)
                st["current_subclause"   ] = new_sub
                st["current_subsubclause"] = None
                st["subsubclause_counter"] = 0
                return (st["current_clause"], False)

        # Subclause detection: single letter like "a." or "(a)"
        m_sub = re.match(r'^\s*\(?([A-Za-z])\)?\s*[\.\)]\s*(.+)$', raw)
        if m_sub and st["current_clause"] is not None:
            letter = m_sub.group(1).lower()
            body = sanitize_text(m_sub.group(2).strip())
            if letter.isalpha() and len(letter) == 1:
                idx = ord(letter) - ord('a') + 1
            else:
                st["subclause_counter"] += 1
                idx = st["subclause_counter"]

            new_sub = subclause(idx, text=body)
            st["current_clause"      ].listsubclauses.append(new_sub)
            st["current_subclause"   ] = new_sub
            st["current_subsubclause"] = None
            st["subsubclause_counter"] = 0
            return (st["current_clause"], False)

        # Continuation: append to most recent item
        cont = raw
        if st["current_subsubclause"] is not None:
            combined = sanitize_text(strip_punctuations(st["current_subsubclause"].text + " " + cont))
            st["current_subsubclause"].text = combined
            return (st["current_clause"], False)
        if st["current_subclause"] is not None:
            combined = sanitize_text(strip_punctuations(st["current_subclause"].text + " " + cont))
            st["current_subclause"].text = combined
            return (st["current_clause"], False)
        if st["current_clause"] is not None:
            combined = sanitize_text(strip_punctuations(st["current_clause"].text + " " + cont))
            st["current_clause"].text = combined
            return (st["current_clause"], False)

        # nothing matched
        return (clause(0, "__ERROR__", raw), False)

    components['preambs'     ] = cast(_rc_t, ResolutionComponent[preamb]())
    components['operationals'] = cast(_rc_t, ResolutionComponent[clause]())

    components['preambs'     ].matchFunc = _preambs_match_function
    components['operationals'].matchFunc = _operationals_match_function
    componentsList: list[str] = ['committee',
                                 'mainSubmitter',
                                 'coSubmitters',
                                 'topic',
                                 'preambs',
                                 'operationals']

    # ====== Main Loop ======
    # ====== Main Loop ======
    for index, line in enumerate(paragraphs):
        text = line.strip()
        if verbose:
            print(f"{Fore.MAGENTA}{index:3}{Style.RESET_ALL}| {line}")

        for componentName in componentsList:
            comp = components[componentName]

            if comp.matchFunc is not None:
                try:
                    val, ok = comp.matchFunc(text)
                    if val is not None:
                        comp.appendValue(val)
                    if not ok and val is not None and getattr(val, "adverb", "") == "__ERROR__":
                        errorList.append(
                            ResolutionParsingError(
                                f"{componentName} could not parse: {text}",
                                index + 1
                            )
                        )
                except Exception as e:
                    errorList.append(
                        ResolutionParsingError(
                            f"{componentName} parser exception: {e}",
                            index + 1
                        )
                    )

            else:
                # plain regex components just try to extract
                before = len(comp.getValues())
                comp.extract(text, re.IGNORECASE)
                after = len(comp.getValues())
                # don’t add errors here – handle at the end if still empty


    # finalize values (preambs and operationals built by matchFuncs)
    components['preambs'     ].setValue(cast(list[_rc_inner_t], listPreambs))
    components['preambs'     ].markFinished()

    components['operationals'].setValue(cast(list[_rc_inner_t], listOperationals))
    components['operationals'].markFinished()
    
    # Add errors if key header components were never found at all
    for cname in ["committee", "mainSubmitter", "topic"]:
        if not components[cname].getValues():
            errorList.append(
                ResolutionParsingError(f"Missing required field: {cname}", -1)
            )


    reso = Resolution(
        cast(str, components['committee'].getFirst()),
        cast(str, components['mainSubmitter'].getFirst()),
        cast(list[str], dedupe_preserve_order(components['coSubmitters'].getListValues() or ['None'])),
        cast(str, components['topic'].getFirst()),
    )

    reso.preambs = listPreambs
    reso.clauses = listOperationals

    return (reso, components, errorList)

def writeToFile(resolution, filename: str | Path) -> int:
    outDoc = doc.document(None, str(filename), line_spacing=2)

    topicPar = doc.paragraph(bold=True)
    topicPar.add_run("Topic: ", bold=True)
    topicPar.add_run(resolution.topic, bold=False)

    committeePar = doc.paragraph(bold=True)
    committeePar.add_run("Committee: ", bold=True)
    committeePar.add_run(resolution.committee, bold=False)

    mainSubmitterPar = doc.paragraph(bold=True)
    mainSubmitterPar.add_run("Main Submitter: ", bold=True)
    mainSubmitterPar.add_run(resolution.mainSubmitter, bold=False)

    coSubmittersPar = doc.paragraph(bold=True)
    coSubmittersPar.add_run("Co-Submitters: ", bold=True)
    coSubmittersPar.add_run(", ".join(resolution.coSubmitters), bold=False)

    committeeSubjectPar = doc.paragraph(
        f"The {' '.join(
            word.capitalize() for word in resolution.committee.split(' ')[:-1]
        )},",
        bold=False
    )

    # --- preambs ---
    preambs: list[doc.paragraph] = []
    for pre in resolution.preambs:
        temp = doc.paragraph()
        temp.add_run(pre.adverb.capitalize(), italic=True)
        temp.add_run(" " + pre.content + ",", italic=False)
        preambs.append(temp)

    # --- operationals ---

   
    def _ends_with_list_phrase(text: str) -> bool:
        # Return True if text ends with any phrase in list_phrases (robust to trailing punctuation/whitespace)
        if not text:
            return False
        for ph in list_phrases:
            # match phrase at the end, allow optional trailing spaces and punctuation
            pattern = rf'(?i)\b{re.escape(ph)}[ \t\r\n]*[,:;.\-–—]*$'
            if re.search(pattern, text):
                return True
        return False

    def render_clause(theclause: clause | subclause | subsubclause,
                      level: int = 1,
                      is_last: bool = False,
                      last_within_top_clause: bool = False) -> list[doc.paragraph]:
        """
        is_last: absolute last in the whole resolution (-> period).
        last_within_top_clause: this element is the last item inside its top-level clause
            (used to decide whether to use semicolon for leaf items).
        """
        paragraphs: list[doc.paragraph] = []

        def choose_end(text: str, is_last: bool, last_within_top_clause: bool, has_children: bool = False) -> str:
            base = text.rstrip(",.").rstrip(":").rstrip(";")
            # 1) if it has children or ends with a list phrase -> colon (overrides absolute last)
            if has_children or _ends_with_list_phrase(text):
                return base + ":"
            # 2) absolute last leaf -> period
            if is_last:
                return base + "."
            # 3) last leaf within top clause -> semicolon
            if last_within_top_clause:
                return base + ";"
            # 4) otherwise comma
            return base + ","


        # Clause-level formatting
        if isinstance(theclause, clause):
            par = doc.paragraph(list_level=1)
            par.add_run(theclause.verb.capitalize() + " ", underline=True)
            has_children = bool(getattr(theclause, "listsubclauses", []))
            par.add_run(choose_end(theclause.text, is_last, False, has_children))
            paragraphs.append(par)

            # Recurse into subclauses
            for i, sub in enumerate(theclause.listsubclauses):
                # whether this subclause is the last subclause of this clause
                sub_is_last_in_clause = (i == len(theclause.listsubclauses) - 1)
                # absolute last for this sub = absolute last for parent clause AND this is its last sub
                sub_is_absolute_last = is_last and sub_is_last_in_clause
                # last_within_top_clause for the sub = whether it's the last subclause inside the top clause
                paragraphs.extend(render_clause(sub, 2, sub_is_absolute_last, sub_is_last_in_clause))

        elif isinstance(theclause, subclause):
            has_children = bool(getattr(theclause, "listsubsubclauses", []))
            par = doc.paragraph(list_level=level)
            par.add_run(choose_end(theclause.text, is_last, last_within_top_clause, has_children))
            paragraphs.append(par)

            # Recurse into sub-subclauses
            for i, subsub in enumerate(theclause.listsubsubclauses):
                subsub_is_last_in_sub = (i == len(theclause.listsubsubclauses) - 1)
                # absolute last for this subsub = absolute last passed down AND it's the last subsub
                subsub_is_absolute_last = is_last and subsub_is_last_in_sub
                # last_within_top_clause for subsub: it is last within top clause only if:
                #   (a) parent subclause was last_within_top_clause (i.e. parent was last subclause of clause)
                #   AND (b) this subsub is the last in its parent subclause
                subsub_last_within_top_clause = last_within_top_clause and subsub_is_last_in_sub
                paragraphs.extend(render_clause(subsub, 3, subsub_is_absolute_last, subsub_last_within_top_clause))

        elif isinstance(theclause, subsubclause):
            par = doc.paragraph(list_level=level)
            par.add_run(choose_end(theclause.text, is_last, last_within_top_clause))
            paragraphs.append(par)

        return paragraphs


    for par in [
        topicPar, committeePar, mainSubmitterPar,
        coSubmittersPar, committeeSubjectPar
    ]:
        outDoc.append(par)

    for pre in preambs:
        outDoc.append(pre)

    # Render all clauses, passing is_last=True to the last clause
    for i, cl in enumerate(resolution.clauses):
        is_last_clause = (i == len(resolution.clauses) - 1)
        # For top-level clauses, last_within_top_clause is the same as is_last_clause
        pars = render_clause(cl, level=1, is_last=is_last_clause, last_within_top_clause=is_last_clause)
        for par in pars:
            outDoc.append(par)

    outDoc.save(verbose=verbose)
    return 0


def main():
    global verbose
    # Set default filenames
    input_filename: str | Path = Path("../tests/inputs/test_reso.docx")
    output_filename: str | Path = Path("../tests/outputs/test_reso.docx")
    log_filename: str | Path | None = Path("../tests/outputs/formatter.log")

    parser = argparse.ArgumentParser(
                    prog='',
                    description='Formats a resolution (.docx) and outputs file.')
    parser.add_argument('filename', nargs='?', help='input filename (optional)')  # Changed to optional
    parser.add_argument('-v', '--verbose', help='enable verbose mode', action='store_true')
    parser.add_argument('-o', '--output', nargs='?', help='output filename')
    parser.add_argument('-l', '--log', nargs='?', help="log file name")
    args = parser.parse_args()

    if args.verbose:
        verbose = True
    
    # Use the positional filename argument if provided
    if args.filename:
        if verbose:
            print(f"Input filename: {args.filename}")
        input_filename = Path(args.filename)
    
    if args.output:
        if verbose:
            print(f"Output filename: {args.output}")
        output_filename = Path(args.output)
    else:
        print(f"{Fore.RED}Changing the original file as output filename is not given{Style.RESET_ALL}")
        output_filename = input_filename
    if verbose:
        print(f"Using input: {input_filename}")
        print(f"Using output: {output_filename}")
    
    if args.log:
        if verbose:
            print(f"Log filename: {args.log}")
        log_filename = Path(args.log)
    

    """
    Step 1: Read doc and parse to object
    """
    try:
        resolutionRawDocument = doc.document(str(input_filename), str(output_filename))
        parseResult = parseToResolution(resolutionRawDocument)
        parsedResolution, components, errorList = parseResult
    except PackageNotFoundError:
        print(f"{Fore.RED}{Style.BRIGHT}Error: invalid input / output path{Style.RESET_ALL}")
        return

    if verbose: print(str(parsedResolution))


    """
    Step 2: Conflict/Error showing, resolution and confirmation
    """

    if (len(errorList) != 0) and args.log:
        print(f"{Fore.RED}Errors found and corrected. Check {log_filename} for log / errors")
        with open(str(log_filename), "w") as f:
            f.write(f"ERROR LOG FOR {os.path.abspath(input_filename)}\n")
            for error in errorList:
                f.write(str(error) + "\n")
    elif len(errorList) != 0:
        print(f"{Fore.RED}Errors in resolution:{Style.RESET_ALL}")
        for error in errorList:
            print(f"{Fore.MAGENTA}{str(error)}{Style.RESET_ALL}")


    """
    Step 3: Write to file
    """

    writeToFile(parsedResolution, output_filename)

if __name__ == "__main__":
    main()