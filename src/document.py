from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.numbering import CT_Numbering
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from typing import Union


class ResoFormattingError(BaseException):
    def __init__(self, msg: str | None = None):
        super().__init__(msg)



class NumberingStyleManager:
    """
    A multilevel list implementation that:
     1. Builds a new <w:abstractNum> + <w:num> in numbering.xml
     2. Appends them directly to the CT_Numbering root
     3. Reuses that numId for all paragraphs
    """

    def __init__(self, doc):
        self.doc = doc
        self.num_id = None
        self._ensure_numbering()

    def _ensure_numbering(self):
        num_part = self.doc.part.numbering_part
        ct_num = num_part.element  # CT_Numbering element

        # 1) find all existing abstractNumId and numId values
        abs_ids = [int(el.get(qn('w:abstractNumId')))
                   for el in ct_num.findall(qn('w:abstractNum'))]
        num_ids = [int(el.get(qn('w:numId')))
                   for el in ct_num.findall(qn('w:num'))]

        new_abs_id = max(abs_ids + [0]) + 1
        new_num_id = max(num_ids + [0]) + 1
        self.num_id = str(new_num_id)

        # 2) build abstractNum
        abstractNum = OxmlElement('w:abstractNum')
        abstractNum.set(qn('w:abstractNumId'), str(new_abs_id))

        # define up to 4 levels (decimal, a, i, decimal)
        level_formats = {
            0: 'decimal',
            1: 'lowerLetter',
            2: 'lowerRoman',
            3: 'decimal',
        }
        for lvl, fmt in level_formats.items():
            lvl_elm = OxmlElement('w:lvl')
            lvl_elm.set(qn('w:ilvl'), str(lvl))

            start = OxmlElement('w:start')
            start.set(qn('w:val'), '1')
            numFmt = OxmlElement('w:numFmt')
            numFmt.set(qn('w:val'), fmt)
            lvlText = OxmlElement('w:lvlText')
            # %1., %2., etc.
            lvlText.set(qn('w:val'), f'%{lvl+1}.')
            suff = OxmlElement('w:suff')
            suff.set(qn('w:val'), 'space')

            for node in (start, numFmt, lvlText, suff):
                lvl_elm.append(node)
            abstractNum.append(lvl_elm)

        # 3) build concrete <w:num> that references our abstractNum
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), self.num_id)
        abstrId = OxmlElement('w:abstractNumId')
        abstrId.set(qn('w:val'), str(new_abs_id))
        num.append(abstrId)

        # 4) append both to the CT_Numbering element
        ct_num.append(abstractNum)
        ct_num.append(num)

    def add_numbered_paragraph(self, text, level=1):
        # create a normal paragraph, then inject numbering XML
        p = self.doc.add_paragraph()
        p.add_run(text)

        numPr = p._p.get_or_add_pPr().get_or_add_numPr()

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level-1))
        numPr.append(ilvl)

        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), self.num_id)
        numPr.append(numId)

        # indent by 0.5" per level
        p.paragraph_format.left_indent = Inches(0.31988 * level)
        p.paragraph_format.first_line_indent = Inches(-0.31988)

        return p
class paragraph:
    def __init__(
        self,
        text: str = "",
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        font_size: int = 12,
        font_color: tuple[int, int, int] | None = None,
        align: Union[WD_PARAGRAPH_ALIGNMENT, str, None] = None,
        style: str | None = None,
        first_line_indent: Union[float | int, Pt] | None = None, # float | int -> Inches, Pt -> Pt
        left_indent: Union[float | int, Pt] | None = None,       # float | int -> Inches, Pt -> Pt
        right_indent: Union[float | int, Pt] | None = None,      # float | int -> Inches, Pt -> Pt
        list_level: int = 0, # 0: no list, 1: level 1, 2: level 2 (alpha), 3: level 3 (roman)
        list_format: str = 'decimal',
        list_indents: dict[int, float] | None = None,
        line_spacing: float | None = None, # None = inherit from parent document
    ) -> None:
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.font_size = font_size
        self.font_color = font_color
        self.style = style
        self._runs: list[dict] = []
        self.align = None
        self.first_line_indent = first_line_indent
        self.left_indent = left_indent
        self.right_indent = right_indent
        self.list_level = list_level
        self.list_format = list_format
        self.list_indents = list_indents or {}
        self.line_spacing = line_spacing
        
        if align is not None:
            self.set_alignment(align) # type: ignore

    def set_alignment(self, align: Union[WD_PARAGRAPH_ALIGNMENT, str]) -> None: # type: ignore
        """Set paragraph alignment using either enum or string."""
        if isinstance(align, WD_PARAGRAPH_ALIGNMENT):
            self.align = align
        elif isinstance(align, str):
            align_lower = align.lower()
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                "distribute": WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE,
            }
            if align_lower in alignment_map:
                self.align = alignment_map[align_lower]
            else:
                raise ValueError(f"Invalid alignment string: {align}")
        else:
            raise TypeError("Alignment must be WD_PARAGRAPH_ALIGNMENT or str")

    def add_run(self, text: str, **kwargs) -> None:
        """Add a formatted text run.
        kwargs choices: bold, italic, underline, font_size
        """
        self._runs.append({
            "text": text,
            "bold": kwargs.get("bold", False),
            "italic": kwargs.get("italic", False),
            "underline": kwargs.get("underline", False),
            "font_size": kwargs.get("font_size", self.font_size),
        })

    def render(self, doc: Document) -> None: # type: ignore
        """Render the paragraph to a docx Document."""
        # Handle list paragraphs differently
        if self.list_level > 0:
            p = NumberingStyleManager(doc).add_numbered_paragraph("", self.list_level)
            run = p.add_run(self.text)
            self._apply_formatting(run)
        else:
            # Regular paragraph handling
            p = doc.add_paragraph(style=self.style)
            
            # If we have runs, add them with formatting
            if self._runs:
                for run_spec in self._runs:
                    run = p.add_run(run_spec["text"])
                    run.bold = run_spec["bold"]
                    run.italic = run_spec["italic"]
                    run.underline = run_spec["underline"]
                    if run_spec["font_size"]:
                        run.font.size = Pt(run_spec["font_size"])
                    if self.font_color:
                        run.font.color.rgb = RGBColor(*self.font_color)
            else:
                # Add simple text with formatting
                run = p.add_run(self.text)
                self._apply_formatting(run)

        # Apply paragraph formatting (alignment, indents, etc.)
        if self.align is not None:
            p.alignment = self.align

        if self.first_line_indent is not None:
            if isinstance(self.first_line_indent, (int, float)):
                p.paragraph_format.first_line_indent = Inches(self.first_line_indent)
            else:
                p.paragraph_format.first_line_indent = self.first_line_indent
                
        if self.left_indent is not None:
            if isinstance(self.left_indent, (int, float)):
                p.paragraph_format.left_indent = Inches(self.left_indent)
            else:
                p.paragraph_format.left_indent = self.left_indent
                
        if self.right_indent is not None:
            if isinstance(self.right_indent, (int, float)):
                p.paragraph_format.right_indent = Inches(self.right_indent)
            else:
                p.paragraph_format.right_indent = self.right_indent
        
        # line spacing
        p.paragraph_format.line_spacing = self.line_spacing if self.line_spacing is not None else doc.styles['Normal'].paragraph_format.line_spacing # potential bug here if overall style is not 'Normal', but considering our specific use case it can be tolerated (plus nobody formats their entire document in a style other than Normal anyway)


    def _apply_formatting(self, run) -> None:
        """Helper to apply formatting to a run."""
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline
        run.font.size = Pt(self.font_size)
        if self.font_color:
            run.font.color.rgb = RGBColor(*self.font_color)
    
    
    def set_style(self, style: str) -> None:
        """Set paragraph style (e.g., 'Heading 1')"""
        self.style = style
    
    def set_font_color(self, r: int, g: int, b: int) -> None:
        """Set font color using RGB values (0-255)"""
        self.font_color = (r, g, b)
        
    # float | int -> Inches, Pt -> Pt
    def set_first_line_indent(self, value: Union[float | int, Pt]) -> None:
        """Set first line indent (in inches or Pt)"""
        self.first_line_indent = Inches(value) if isinstance(value, (int, float)) else (value)

    # float | int -> Inches, Pt -> Pt
    def set_left_indent(self, value: Union[float | int, Pt]) -> None:
        """Set left indent (in inches or Pt)"""
        self.left_indent = Inches(value) if isinstance(value, (int, float)) else (value)
    
    def __str__(self) -> str:
        """Return plain text content"""
        if self._runs:
            return ''.join(str(r['text']) for r in self._runs)  # Explicitly convert to str
        return self.text
    
    

class document:
    def __init__(self, 
                 inputfile: str ="tests/inputs/test1.docx",
                 outputfile : str ="tests/outputs/test1.docx",
                 paragraphs: list[paragraph] | None = None,
                 overallstyle : str = "Normal",
                 font: str = "Times New Roman",
                 fontsize : int | float = 12,
                 line_spacing: int | float = 1):
        self.paragraphs = paragraphs if paragraphs is not None else []
        self.inputfile = inputfile
        self.outputfile = outputfile
        self._doc = Document(inputfile)
        self._doc.styles[overallstyle].font.name = font # type: ignore
        self._doc.styles[overallstyle].font.size = Pt(fontsize) # type: ignore
        self._doc.styles[overallstyle].paragraph_format.line_spacing = line_spacing # type: ignore
        
    def append(self, paragraph: paragraph, index: int | None = None) -> 'document':
        """
        Append or insert a paragraph to the document.
        
        Args:
            paragraph: The paragraph object to add
            index: Optional index position to insert at. If None, appends to end.
        
        Returns:
            document: Self for method chaining
        """
        if index is None:
            # Append to the end
            self.paragraphs.append(paragraph)
            paragraph.render(self._doc)
        else:
            # Insert at specified index
            self.paragraphs.insert(index, paragraph)
            
            # Clear the current document content and re-render all paragraphs
            self.rebuild_document()
        
        return self

    def rebuild_document(self) -> None:
        """
        Rebuild the document by clearing all content and re-rendering all paragraphs.
        This is necessary when inserting paragraphs at specific positions.
        """
        # Clear all existing paragraphs from the document
        for element in self._doc.element.body:
            if element.tag.endswith('p'):  # Remove paragraph elements
                self._doc.element.body.remove(element)
        
        # Re-render all paragraphs in the correct order
        for para in self.paragraphs:
            para.render(self._doc)
    
    def remove(self, paragraph: paragraph) -> None:
        self.paragraphs.remove(paragraph)
    def save(self, outputfile : str | None = None) -> None:
        if outputfile is None:
            outputfile = self.outputfile
        self._doc.save(outputfile)
        print(f"File saved to {outputfile}")
    def getdocument(self) -> 'Document': # type: ignore
        return self._doc

    def get_paragraphs(self) -> list[str]:
        """
        Returns a list of text content from all paragraphs in the document,
        including proper hierarchical numbering information.
        """
        paragraphs_text = []
        # Track numbering state for each list and level
        numbering_states = {}
        current_list_context = None
        
        for i, paragraph in enumerate(self._doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Extract numbering information with proper hierarchy
            numbering_text, list_context = self._extract_hierarchical_numbering(
                paragraph, numbering_states, current_list_context
            )
            
            if list_context:
                current_list_context = list_context
                
            if numbering_text:
                paragraphs_text.append(f"{numbering_text} {text}")
            else:
                paragraphs_text.append(text)
        
        return paragraphs_text

    def _extract_hierarchical_numbering(self, 
                                        paragraph,
                                        numbering_states, 
                                        current_list_context) -> tuple:
        """
        Extract numbering information from a paragraph - show only current level number.
        """
        try:
            pPr = paragraph._p.pPr
            if pPr is None:
                return "", current_list_context
                
            numPr = pPr.numPr
            if numPr is None:
                return "", current_list_context
                
            # Get numbering ID and level
            numId_elem = numPr.numId
            ilvl_elem = numPr.ilvl
            
            if numId_elem is None or ilvl_elem is None:
                return "", current_list_context
                
            numId = numId_elem.val
            ilvl = ilvl_elem.val
            
            if numId is None or ilvl is None:
                return "", current_list_context
                
            # Get numbering definition
            numbering_part = self._doc.part.numbering_part
            if numbering_part is None:
                return "", current_list_context
                
            # Find the abstract numbering definition
            abstract_num_id = self._get_abstract_num_id(numbering_part, numId)
            if abstract_num_id is None:
                return "", current_list_context
                
            # Create a key for this numbering context
            list_context = (numId, abstract_num_id)
            
            # Initialize numbering state if it doesn't exist
            if list_context not in numbering_states:
                numbering_states[list_context] = {}
                
            # Get numbering format for this level
            num_format = self._get_number_format(numbering_part, abstract_num_id, ilvl)
            if not num_format:
                return "", list_context
                
            # Update numbering state
            level_state = numbering_states[list_context]
            
            # Reset higher levels if we're going back to a lower level
            for l in range(ilvl + 1, 10):  # Assuming max 10 levels
                if l in level_state:
                    del level_state[l]
                    
            # Initialize or increment current level
            if ilvl not in level_state:
                level_state[ilvl] = 0
            level_state[ilvl] += 1
            
            # Return ONLY the current level's number (not the full hierarchy)
            current_number = self._format_number(level_state[ilvl], num_format)
            
            return current_number + ".", list_context
            
        except (AttributeError, TypeError, KeyError):
            return "", current_list_context

    def _continue_current_numbering(self, list_context, numbering_states):
        """
        Continue the current numbering without incrementing the level.
        Returns only the current level's number.
        """
        if list_context not in numbering_states:
            return ""
            
        level_state = numbering_states[list_context]
        if not level_state:
            return ""
            
        # Find the highest level with a value (current level)
        if not level_state:
            return ""
            
        current_level = max(level_state.keys())
        current_number = level_state[current_level]
        
        # Get numbering format for current level
        numbering_part = self._doc.part.numbering_part
        if numbering_part is None:
            return ""
            
        abstract_num_id = list_context[1]
        num_format = self._get_number_format(numbering_part, abstract_num_id, current_level)
        
        # Format only the current level's number
        formatted_number = self._format_number(current_number, num_format or "decimal")
        
        return formatted_number + "."

    def _get_abstract_num_id(self, numbering_part, numId):
        """Get the abstract numbering ID for a concrete numbering ID."""
        for num in numbering_part.element.findall(qn('w:num')):
            if num.get(qn('w:numId')) == str(numId):
                abstractNumId = num.find(qn('w:abstractNumId'))
                if abstractNumId is not None:
                    return abstractNumId.get(qn('w:val'))
        return None

    def _get_number_format(self, numbering_part, abstract_num_id, ilvl):
        """Get the number format for a specific level in a numbering definition."""
        for abstractNum in numbering_part.element.findall(qn('w:abstractNum')):
            if abstractNum.get(qn('w:abstractNumId')) == abstract_num_id:
                for lvl in abstractNum.findall(qn('w:lvl')):
                    if lvl.get(qn('w:ilvl')) == str(ilvl):
                        numFmt = lvl.find(qn('w:numFmt'))
                        if numFmt is not None:
                            return numFmt.get(qn('w:val'))
        return None

    def _format_number(self, number, num_format):
        """Format a number according to the specified format."""
        format_map = {
            'decimal': str(number),
            'lowerLetter': self._number_to_letters(number).lower(),
            'upperLetter': self._number_to_letters(number),
            'lowerRoman': self._number_to_roman(number).lower(),
            'upperRoman': self._number_to_roman(number),
            'bullet': '•',
        }
        
        return format_map.get(num_format, str(number))

    def _number_to_letters(self, num):
        """Convert a number to letters (1 = A, 2 = B, ..., 27 = AA, etc.)."""
        letters = ''
        while num > 0:
            num, remainder = divmod(num - 1, 26)
            letters = chr(65 + remainder) + letters
        return letters

    def _number_to_roman(self, num):
        """Convert a number to Roman numerals."""
        val = [
            1000, 900, 500, 400,
            100, 90, 50, 40,
            10, 9, 5, 4,
            1
        ]
        syb = [
            "M", "CM", "D", "CD",
            "C", "XC", "L", "XL",
            "X", "IX", "V", "IV",
            "I"
        ]
        roman_num = ''
        i = 0
        while num > 0:
            for _ in range(num // val[i]):
                roman_num += syb[i]
                num -= val[i]
            i += 1
        return roman_num